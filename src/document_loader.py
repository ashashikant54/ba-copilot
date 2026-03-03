# document_loader.py
# UPDATED FOR AZURE CLOUD DEPLOYMENT
#
# WHAT CHANGED FROM THE ORIGINAL:
#   - Added import: io  (built-in Python library, no install needed)
#   - Added 3 new functions for reading from bytes (cloud uploads):
#       load_txt_from_bytes()
#       load_docx_from_bytes()
#       load_pdf_from_bytes()
#   - Added 1 new main function:
#       load_document_from_bytes()  ← use this in Azure instead of load_document()
#
# WHAT STAYED THE SAME:
#   - load_txt(), load_docx(), load_pdf(), load_document() are ALL UNCHANGED
#   - These still work perfectly for local testing on your laptop
#
# WHY WE ADDED INSTEAD OF REPLACED:
#   - On your LAPTOP: files are saved to disk → use load_document(file_path)
#   - On AZURE: files arrive as bytes in memory → use load_document_from_bytes()
#   - Azure App Service's disk is wiped on restart, so we never save uploaded
#     files to disk — we read them directly from the upload request in memory

import os
import io   # NEW: built-in Python library for working with data in memory
            # No need to pip install anything — it's built into Python


# ── ORIGINAL FUNCTIONS (COMPLETELY UNCHANGED) ─────────────────
# These still work exactly as before for local development

def load_txt(file_path):
    """Read a plain text file"""
    with open(file_path, "r", encoding="utf-8") as f:
        text = f.read()
    return text


def load_docx(file_path):
    """Read a Word document (.docx)"""
    from docx import Document
    doc = Document(file_path)
    paragraphs = []
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            paragraphs.append(paragraph.text)
    return "\n".join(paragraphs)


def load_pdf(file_path):
    """Read a PDF file"""
    from pypdf import PdfReader
    reader = PdfReader(file_path)
    pages = []
    for page in reader.pages:
        text = page.extract_text()
        if text:
            pages.append(text)
    return "\n".join(pages)


def load_document(file_path):
    """
    ORIGINAL: Load from a file path on disk.
    Still works for local testing on your laptop.
    On Azure, use load_document_from_bytes() instead.
    """
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"Cannot find file: {file_path}")

    extension = os.path.splitext(file_path)[1].lower()

    if extension == ".txt":
        text = load_txt(file_path)
    elif extension == ".docx":
        text = load_docx(file_path)
    elif extension == ".pdf":
        text = load_pdf(file_path)
    else:
        raise ValueError(f"Sorry, I can only read .txt, .docx and .pdf files. You gave me: {extension}")

    print(f"✅ Loaded: {file_path}")
    print(f"📄 Characters read: {len(text)}")
    print(f"📝 Words read: {len(text.split())}")

    return text


# ── NEW FUNCTIONS FOR AZURE (CLOUD UPLOADS) ───────────────────
# These do the SAME job as the functions above, but work with
# bytes in memory instead of files on disk.
#
# What is "bytes"?
#   When a user uploads a file through your web UI, FastAPI receives
#   it as raw bytes — basically the file's content held in memory.
#   io.BytesIO() wraps those bytes so libraries like pypdf and docx
#   can read them exactly as if they were a normal file.

def load_txt_from_bytes(file_bytes):
    """
    NEW: Read a plain text file from bytes (cloud upload).
    file_bytes = the raw content of the file, received from the upload
    """
    # Decode bytes → string, same as reading a text file
    return file_bytes.decode("utf-8")


def load_docx_from_bytes(file_bytes):
    """
    NEW: Read a Word document from bytes (cloud upload).
    io.BytesIO() makes bytes look like a file to the docx library.
    """
    from docx import Document
    # io.BytesIO() = "pretend these bytes are a file on disk"
    doc = Document(io.BytesIO(file_bytes))
    paragraphs = []
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            paragraphs.append(paragraph.text)
    return "\n".join(paragraphs)


def load_pdf_from_bytes(file_bytes):
    """
    NEW: Read a PDF file from bytes (cloud upload).
    io.BytesIO() makes bytes look like a file to the pypdf library.
    """
    from pypdf import PdfReader
    # io.BytesIO() = "pretend these bytes are a file on disk"
    reader = PdfReader(io.BytesIO(file_bytes))
    pages = []
    for page in reader.pages:
        text = page.extract_text()
        if text:
            pages.append(text)
    return "\n".join(pages)


def load_document_from_bytes(filename, file_bytes):
    """
    NEW MAIN FUNCTION FOR AZURE: Load a document from bytes.
    
    Use this in your FastAPI upload endpoint instead of load_document().
    
    How to call it in main.py:
        # FastAPI gives you the file like this:
        contents = await file.read()           # file_bytes
        text = load_document_from_bytes(file.filename, contents)
    
    Arguments:
        filename   = the original file name e.g. "HR_Policy.pdf"
                     (we need this just to detect the file type)
        file_bytes = the raw file content received from the upload
    """
    # Get the file extension from the filename (same logic as before)
    extension = os.path.splitext(filename)[1].lower()

    if extension == ".txt":
        text = load_txt_from_bytes(file_bytes)
    elif extension == ".docx":
        text = load_docx_from_bytes(file_bytes)
    elif extension == ".pdf":
        text = load_pdf_from_bytes(file_bytes)
    else:
        raise ValueError(
            f"Sorry, I can only read .txt, .docx and .pdf files. "
            f"You gave me: {extension}"
        )

    print(f"✅ Loaded from upload: {filename}")
    print(f"📄 Characters read: {len(text)}")
    print(f"📝 Words read: {len(text.split())}")

    return text


# ── TEST ──────────────────────────────────────────────────────
if __name__ == "__main__":
    print("Testing document loader...\n")

    # ── Test 1: Original local file loading (unchanged) ──
    file = "documents/sample_requirements.txt"
    if os.path.exists(file):
        print("── Test 1: Local file loading ──")
        text = load_document(file)
        print("\n── First 300 characters ──")
        print(text[:300])
        print("...")

    # ── Test 2: New bytes loading (simulates cloud upload) ──
    print("\n── Test 2: Bytes loading (simulates Azure upload) ──")
    sample_text = "This is a test document.\nIt has two lines."
    fake_upload_bytes = sample_text.encode("utf-8")  # Convert string → bytes
    result = load_document_from_bytes("test_doc.txt", fake_upload_bytes)
    print(f"Result: {result}")

    print("\n✅ Document loader is working!")